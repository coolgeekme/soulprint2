#!/usr/bin/env python3
"""
Backend API Testing - PDF/DOCX Document Parsing API
Tests the POST /api/parse/document endpoint as per review request
"""

import requests
import json
import os
import sys
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
import io

# Base URL from environment
BASE_URL = os.getenv('NEXT_PUBLIC_BASE_URL', 'https://doc-parser-16.preview.emergentagent.com')
API_BASE = f"{BASE_URL}/api"

def create_test_pdf():
    """Create a simple test PDF file for testing."""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        
        content = []
        
        # Add title
        title = Paragraph("Test PDF Document", styles['Title'])
        content.append(title)
        content.append(Spacer(1, 12))
        
        # Add paragraphs
        para1 = Paragraph("This is a test PDF document created for API testing purposes.", styles['Normal'])
        content.append(para1)
        content.append(Spacer(1, 12))
        
        para2 = Paragraph("It contains multiple paragraphs with different content to test text extraction capabilities.", styles['Normal'])
        content.append(para2)
        content.append(Spacer(1, 12))
        
        para3 = Paragraph("The PDF parsing should extract all this text correctly and return it in the API response.", styles['Normal'])
        content.append(para3)
        content.append(Spacer(1, 12))
        
        para4 = Paragraph("This helps verify that the pdf-parse library integration is working properly.", styles['Normal'])
        content.append(para4)
        
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Error creating PDF: {e}")
        # Fallback: create minimal PDF with reportlab basic canvas
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "Test PDF Document")
        c.drawString(100, 730, "This is a simple test PDF for API testing.")
        c.drawString(100, 710, "It should be parsed correctly by the document parsing API.")
        c.drawString(100, 690, "The pdf-parse library should extract this text content.")
        c.save()
        buffer.seek(0)
        return buffer.getvalue()

def create_test_docx():
    """Create a simple test DOCX file for testing."""
    try:
        doc = Document()
        doc.add_heading('Test DOCX Document', 0)
        
        doc.add_paragraph('This is a test DOCX document created for API testing purposes.')
        doc.add_paragraph('It contains multiple paragraphs with different content to test text extraction capabilities.')
        doc.add_paragraph('The DOCX parsing should extract all this text correctly using the mammoth library.')
        doc.add_paragraph('This helps verify that the mammoth library integration is working properly for Word documents.')
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        print(f"Error creating DOCX: {e}")
        return None

def create_test_txt():
    """Create a simple test text file for testing."""
    content = """This is a test text file for API testing.
It contains multiple lines of text content.
The plain text parsing should handle this correctly.
This helps verify the text file fallback functionality."""
    return content.encode('utf-8')

def login_and_get_token():
    """Login with test credentials and get authentication token."""
    try:
        print("🔐 STEP 1: Authentication Test")
        
        login_url = f"{API_BASE}/auth/login"
        login_data = {
            "email": "test@soulprint.com",
            "passcode": "test123"
        }
        
        print(f"POST {login_url}")
        print(f"Request body: {json.dumps(login_data, indent=2)}")
        
        response = requests.post(login_url, json=login_data)
        print(f"Status: {response.status_code}")
        print(f"Response: {json.dumps(response.json(), indent=2)}")
        
        if response.status_code == 200:
            data = response.json()
            if 'token' in data:
                print("✅ Authentication successful!")
                return data['token']
            else:
                print("❌ No token in response")
                return None
        else:
            print(f"❌ Authentication failed: {response.text}")
            return None
            
    except Exception as e:
        print(f"❌ Authentication error: {e}")
        return None

def test_pdf_parsing(token):
    """Test PDF document parsing."""
    try:
        print("\n📄 STEP 2: PDF Parsing Test")
        
        # Create test PDF
        pdf_content = create_test_pdf()
        print(f"Created test PDF: {len(pdf_content)} bytes")
        
        # Prepare multipart request
        url = f"{API_BASE}/parse/document"
        headers = {
            'Authorization': f'Bearer {token}'
        }
        files = {
            'file': ('test.pdf', pdf_content, 'application/pdf')
        }
        
        print(f"POST {url}")
        print(f"Headers: Authorization: Bearer {token[:20]}...")
        print(f"Files: test.pdf ({len(pdf_content)} bytes, application/pdf)")
        
        response = requests.post(url, headers=headers, files=files)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            
            # Validate response structure
            if data.get('success') == True:
                if 'text' in data and 'metadata' in data and 'fileName' in data and 'fileType' in data:
                    print("✅ PDF parsing successful!")
                    print(f"✅ Text extracted: {len(data.get('text', ''))} characters")
                    print(f"✅ Metadata present: {data.get('metadata', {})}")
                    print(f"✅ File name: {data.get('fileName')}")
                    print(f"✅ File type: {data.get('fileType')}")
                    return True
                else:
                    print("❌ Missing required response fields")
                    return False
            else:
                print(f"❌ Parsing failed: {data.get('error', 'Unknown error')}")
                return False
        else:
            print(f"❌ Request failed: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ PDF parsing error: {e}")
        return False

def test_docx_parsing(token):
    """Test DOCX document parsing."""
    try:
        print("\n📝 STEP 3: DOCX Parsing Test")
        
        # Create test DOCX
        docx_content = create_test_docx()
        if not docx_content:
            print("❌ Failed to create test DOCX")
            return False
            
        print(f"Created test DOCX: {len(docx_content)} bytes")
        
        # Prepare multipart request
        url = f"{API_BASE}/parse/document"
        headers = {
            'Authorization': f'Bearer {token}'
        }
        files = {
            'file': ('test.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        print(f"POST {url}")
        print(f"Headers: Authorization: Bearer {token[:20]}...")
        print(f"Files: test.docx ({len(docx_content)} bytes, application/vnd.openxmlformats-officedocument.wordprocessingml.document)")
        
        response = requests.post(url, headers=headers, files=files)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            
            # Validate response structure
            if data.get('success') == True:
                if 'text' in data and 'metadata' in data and 'fileName' in data and 'fileType' in data:
                    print("✅ DOCX parsing successful!")
                    print(f"✅ Text extracted: {len(data.get('text', ''))} characters")
                    print(f"✅ Metadata present: {data.get('metadata', {})}")
                    print(f"✅ File name: {data.get('fileName')}")
                    print(f"✅ File type: {data.get('fileType')}")
                    return True
                else:
                    print("❌ Missing required response fields")
                    return False
            else:
                print(f"❌ Parsing failed: {data.get('error', 'Unknown error')}")
                return False
        else:
            print(f"❌ Request failed: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ DOCX parsing error: {e}")
        return False

def test_txt_parsing(token):
    """Test plain text file parsing."""
    try:
        print("\n📄 STEP 4: Plain Text Parsing Test")
        
        # Create test text file
        txt_content = create_test_txt()
        print(f"Created test TXT: {len(txt_content)} bytes")
        
        # Prepare multipart request
        url = f"{API_BASE}/parse/document"
        headers = {
            'Authorization': f'Bearer {token}'
        }
        files = {
            'file': ('test.txt', txt_content, 'text/plain')
        }
        
        print(f"POST {url}")
        print(f"Headers: Authorization: Bearer {token[:20]}...")
        print(f"Files: test.txt ({len(txt_content)} bytes, text/plain)")
        
        response = requests.post(url, headers=headers, files=files)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            
            # Validate response structure
            if data.get('success') == True:
                if 'text' in data and 'metadata' in data and 'fileName' in data and 'fileType' in data:
                    print("✅ Text parsing successful!")
                    print(f"✅ Text extracted: {len(data.get('text', ''))} characters")
                    print(f"✅ Metadata present: {data.get('metadata', {})}")
                    print(f"✅ File name: {data.get('fileName')}")
                    print(f"✅ File type: {data.get('fileType')}")
                    return True
                else:
                    print("❌ Missing required response fields")
                    return False
            else:
                print(f"❌ Parsing failed: {data.get('error', 'Unknown error')}")
                return False
        else:
            print(f"❌ Request failed: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Text parsing error: {e}")
        return False

def test_no_file_error(token):
    """Test error handling when no file is provided."""
    try:
        print("\n⚠️ STEP 5: No File Error Handling Test")
        
        # Prepare request without file
        url = f"{API_BASE}/parse/document"
        headers = {
            'Authorization': f'Bearer {token}',
            'Content-Type': 'multipart/form-data'
        }
        
        print(f"POST {url}")
        print(f"Headers: Authorization: Bearer {token[:20]}...")
        print("Files: (none - testing error handling)")
        
        response = requests.post(url, headers=headers)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 400:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            
            if 'error' in data and 'no file' in data['error'].lower():
                print("✅ Proper error handling for missing file!")
                return True
            else:
                print(f"❌ Unexpected error message: {data.get('error')}")
                return False
        else:
            print(f"❌ Expected 400 status, got {response.status_code}: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ No file error test failed: {e}")
        return False

def test_unsupported_format(token):
    """Test error handling for unsupported file formats."""
    try:
        print("\n🚫 STEP 6: Unsupported Format Error Handling Test")
        
        # Create fake image file (unsupported format)
        fake_image = b"fake_image_content_for_testing"
        
        # Prepare multipart request with unsupported format
        url = f"{API_BASE}/parse/document"
        headers = {
            'Authorization': f'Bearer {token}'
        }
        files = {
            'file': ('test.png', fake_image, 'image/png')
        }
        
        print(f"POST {url}")
        print(f"Headers: Authorization: Bearer {token[:20]}...")
        print(f"Files: test.png ({len(fake_image)} bytes, image/png)")
        
        response = requests.post(url, headers=headers, files=files)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 400:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            
            if 'error' in data and ('unsupported' in data['error'].lower() or 'not supported' in data['error'].lower()):
                print("✅ Proper error handling for unsupported format!")
                return True
            else:
                print(f"❌ Unexpected error message: {data.get('error')}")
                return False
        else:
            print(f"❌ Expected 400 status, got {response.status_code}: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Unsupported format error test failed: {e}")
        return False

def test_authentication_required():
    """Test that authentication is required."""
    try:
        print("\n🔒 STEP 7: Authentication Required Test")
        
        # Create simple test file
        txt_content = create_test_txt()
        
        # Prepare multipart request WITHOUT authentication
        url = f"{API_BASE}/parse/document"
        files = {
            'file': ('test.txt', txt_content, 'text/plain')
        }
        
        print(f"POST {url}")
        print("Headers: (no authorization header)")
        print(f"Files: test.txt ({len(txt_content)} bytes, text/plain)")
        
        response = requests.post(url, files=files)
        print(f"Status: {response.status_code}")
        
        if response.status_code == 401:
            data = response.json()
            print(f"Response: {json.dumps(data, indent=2)}")
            print("✅ Proper authentication requirement!")
            return True
        else:
            print(f"❌ Expected 401 status, got {response.status_code}: {response.text}")
            return False
            
    except Exception as e:
        print(f"❌ Authentication required test failed: {e}")
        return False

def main():
    """Run all PDF/DOCX Document Parsing API tests."""
    print("🧪 PDF/DOCX Document Parsing API Testing Suite")
    print("=" * 60)
    print(f"Base URL: {BASE_URL}")
    print(f"API Base: {API_BASE}")
    print("=" * 60)
    
    results = {
        'authentication': False,
        'pdf_parsing': False,
        'docx_parsing': False,
        'txt_parsing': False,
        'no_file_error': False,
        'unsupported_format': False,
        'auth_required': False
    }
    
    try:
        # Test 1: Authentication
        token = login_and_get_token()
        if token:
            results['authentication'] = True
            
            # Test 2-6: Document parsing tests (require authentication)
            if test_pdf_parsing(token):
                results['pdf_parsing'] = True
                
            if test_docx_parsing(token):
                results['docx_parsing'] = True
                
            if test_txt_parsing(token):
                results['txt_parsing'] = True
                
            if test_no_file_error(token):
                results['no_file_error'] = True
                
            if test_unsupported_format(token):
                results['unsupported_format'] = True
        
        # Test 7: Authentication required (no token)
        if test_authentication_required():
            results['auth_required'] = True
        
    except Exception as e:
        print(f"❌ Testing suite error: {e}")
    
    # Print summary
    print("\n" + "=" * 60)
    print("🏁 TEST SUMMARY")
    print("=" * 60)
    
    total_tests = len(results)
    passed_tests = sum(results.values())
    
    for test_name, passed in results.items():
        status = "✅ PASS" if passed else "❌ FAIL"
        print(f"{test_name.replace('_', ' ').title():<30} {status}")
    
    print("-" * 60)
    print(f"Tests Passed: {passed_tests}/{total_tests}")
    print(f"Success Rate: {(passed_tests/total_tests)*100:.1f}%")
    
    if passed_tests == total_tests:
        print("\n🎉 ALL TESTS PASSED! PDF/DOCX Document Parsing API is working correctly!")
        return True
    else:
        print(f"\n⚠️ {total_tests - passed_tests} test(s) failed. See details above.")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)